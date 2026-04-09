"""
DocumentExtractor
-----------------
Extracts text from PDF, DOCX, and image files using:
  - pdfplumber  (native PDF text layer)
  - pdf2image + pytesseract  (scanned / image-based PDFs)
  - python-docx  (Word documents)
  - Pillow + pytesseract  (PNG / JPG / TIFF / BMP images)
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Tuple

from PIL import Image, ImageEnhance, ImageFilter

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from PDF, DOCX, and image files."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}

    def __init__(self, ocr_lang: str = "eng", enhance: bool = True):
        """
        Parameters
        ----------
        ocr_lang : Tesseract language code, e.g. "eng", "hin", "spa"
        enhance  : Whether to pre-process images before OCR
        """
        self.ocr_lang = ocr_lang
        self.enhance = enhance

    # ── Public entry point ────────────────────────────────────────────────────

    def extract(self, file_path: str) -> Tuple[str, int]:
        """
        Extract all text from *file_path*.

        Returns
        -------
        (text, page_count)
            text       — full extracted text
            page_count — number of pages or sections processed
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")

        if ext == ".pdf":
            return self._extract_pdf(path)
        elif ext == ".docx":
            return self._extract_docx(path)
        else:
            return self._extract_image(path)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> Tuple[str, int]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Install pdfplumber: pip install pdfplumber")

        pages_text: list[str] = []

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                native_text = page.extract_text() or ""
                if native_text.strip():
                    pages_text.append(native_text)
                else:
                    # Fallback: render page as image and OCR it
                    pil_img = page.to_image(resolution=200).original
                    ocr_text = self._ocr_image(pil_img)
                    if ocr_text.strip():
                        pages_text.append(ocr_text)

        full_text = "\n\n--- Page Break ---\n\n".join(pages_text)
        return full_text, len(pages_text)

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, path: Path) -> Tuple[str, int]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = Document(str(path))
        sections: list[str] = []
        current_section: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Use heading style as section delimiter
            if para.style and "Heading" in (para.style.name or ""):
                if current_section:
                    sections.append("\n".join(current_section))
                current_section = [text]
            else:
                current_section.append(text)

        if current_section:
            sections.append("\n".join(current_section))

        # Also extract text from tables
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_texts.append(" | ".join(row_data))

        if table_texts:
            sections.append("\n--- Tables ---\n" + "\n".join(table_texts))

        full_text = "\n\n".join(sections) if sections else ""
        return full_text, max(len(sections), 1)

    # ── Images ────────────────────────────────────────────────────────────────

    def _extract_image(self, path: Path) -> Tuple[str, int]:
        img = Image.open(str(path))
        text = self._ocr_image(img)
        return text, 1

    # ── OCR core ──────────────────────────────────────────────────────────────

    def _ocr_image(self, img: Image.Image) -> str:
        try:
            import pytesseract
        except ImportError:
            raise ImportError("Install pytesseract: pip install pytesseract")

        if self.enhance:
            img = self._preprocess(img)

        config = f"--oem 3 --psm 6 -l {self.ocr_lang}"
        return pytesseract.image_to_string(img, config=config)

    def _preprocess(self, img: Image.Image) -> Image.Image:
        """Improve image quality before OCR."""
        # Convert to RGB if needed
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # Resize if too small (boosts OCR accuracy)
        w, h = img.size
        if w < 1000 or h < 1000:
            scale = max(1000 / w, 1000 / h)
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        # Convert to grayscale
        img = img.convert("L")

        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)

        # Contrast enhancement
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)

        return img
